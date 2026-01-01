"""
DOC/DOCX 파일 파서

- DOCX: python-docx 라이브러리 사용
- DOC: antiword 또는 LibreOffice 변환
"""

import io
import os
import tempfile
import logging
from typing import Optional
from dataclasses import dataclass

from utils.subprocess_utils import (
    run_libreoffice_convert,
    run_antiword,
    LIBREOFFICE_TIMEOUT,
)

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


@dataclass
class DOCXParseResult:
    success: bool
    text: str
    method: str  # "python-docx", "antiword", "libreoffice"
    page_count: int
    error_message: Optional[str] = None


class DOCXParser:
    """DOC/DOCX 파일 파서"""

    # 페이지당 평균 문자 수 (페이지 수 추정용)
    CHARS_PER_PAGE = 2000

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self):
        """의존성 체크"""
        if Document is None:
            logger.warning("python-docx not installed - DOCX parsing will be limited")

    def parse(self, file_bytes: bytes, filename: str = "document.docx") -> DOCXParseResult:
        """
        DOC/DOCX 파일 파싱

        Args:
            file_bytes: 파일 바이트
            filename: 파일명 (확장자 확인용)

        Returns:
            DOCXParseResult 객체
        """
        ext = filename.lower().split('.')[-1] if '.' in filename else ''

        if ext == 'docx':
            return self._parse_docx(file_bytes)
        elif ext == 'doc':
            return self._parse_doc(file_bytes, filename)
        else:
            # 확장자 없으면 DOCX 먼저 시도
            result = self._parse_docx(file_bytes)
            if result.success:
                return result
            return self._parse_doc(file_bytes, filename)

    def _parse_docx(self, file_bytes: bytes) -> DOCXParseResult:
        """
        DOCX 파일 파싱 (python-docx 사용)
        """
        if Document is None:
            return DOCXParseResult(
                success=False,
                text="",
                method="none",
                page_count=0,
                error_message="python-docx not installed"
            )

        try:
            doc = Document(io.BytesIO(file_bytes))

            # 텍스트 추출
            texts = []

            # 본문 단락
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

            # 테이블 내용
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        texts.append(' | '.join(row_texts))

            combined_text = '\n'.join(texts)

            # 페이지 수 추정
            page_count = max(1, len(combined_text) // self.CHARS_PER_PAGE)

            if len(combined_text.strip()) > 0:
                return DOCXParseResult(
                    success=True,
                    text=combined_text,
                    method="python-docx",
                    page_count=page_count
                )
            else:
                return DOCXParseResult(
                    success=False,
                    text="",
                    method="python-docx",
                    page_count=0,
                    error_message="DOCX 파일에서 텍스트를 추출할 수 없습니다."
                )

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return DOCXParseResult(
                success=False,
                text="",
                method="none",
                page_count=0,
                error_message=f"DOCX_PARSE_ERROR: {str(e)}"
            )

    def _parse_doc(self, file_bytes: bytes, filename: str) -> DOCXParseResult:
        """
        DOC 파일 파싱

        1차: antiword 사용
        2차: LibreOffice로 변환
        """
        # 1차: antiword 시도
        try:
            result = self._parse_via_antiword(file_bytes)
            if result.success:
                return result
        except Exception as e:
            logger.warning(f"antiword parsing failed: {e}")

        # 2차: LibreOffice 변환
        try:
            result = self._parse_via_libreoffice(file_bytes, filename)
            if result.success:
                return result
        except Exception as e:
            logger.warning(f"LibreOffice parsing failed: {e}")

        return DOCXParseResult(
            success=False,
            text="",
            method="none",
            page_count=0,
            error_message="DOC_PARSE_FAILED: DOC 파일을 읽을 수 없습니다. DOCX 또는 PDF로 변환 후 업로드해주세요."
        )

    def _parse_via_antiword(self, file_bytes: bytes) -> DOCXParseResult:
        """
        antiword를 사용한 DOC 파싱

        Requirements:
        - antiword 설치 필요 (apt install antiword)

        강화된 타임아웃 및 프로세스 관리 적용
        """
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as f:
            f.write(file_bytes)
            temp_path = f.name

        try:
            result = run_antiword(temp_path, timeout=30)

            if not result.success:
                if result.timed_out:
                    raise Exception("antiword timed out")
                raise Exception(f"antiword failed: {result.error_message}")

            text = result.stdout

            if len(text.strip()) > 0:
                page_count = max(1, len(text) // self.CHARS_PER_PAGE)
                return DOCXParseResult(
                    success=True,
                    text=text,
                    method="antiword",
                    page_count=page_count
                )

            return DOCXParseResult(
                success=False,
                text="",
                method="antiword",
                page_count=0,
                error_message="antiword returned empty text"
            )

        finally:
            try:
                os.unlink(temp_path)
            except Exception:
                pass

    def _parse_via_libreoffice(self, file_bytes: bytes, filename: str) -> DOCXParseResult:
        """
        LibreOffice로 DOC -> DOCX 변환 후 파싱

        Requirements:
        - LibreOffice 설치 필요

        강화된 타임아웃 및 프로세스 관리 적용
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            # DOC 파일 저장
            doc_path = os.path.join(temp_dir, filename)
            with open(doc_path, 'wb') as f:
                f.write(file_bytes)

            # LibreOffice로 DOCX 변환 (강화된 타임아웃)
            conv_result = run_libreoffice_convert(
                input_path=doc_path,
                output_dir=temp_dir,
                output_format="docx",
                timeout=LIBREOFFICE_TIMEOUT
            )

            if not conv_result.success:
                if conv_result.timed_out:
                    raise Exception(f"LibreOffice conversion timed out after {LIBREOFFICE_TIMEOUT}s")
                raise Exception(f"LibreOffice conversion failed: {conv_result.error_message}")

            # 변환된 DOCX 찾기
            docx_filename = os.path.splitext(filename)[0] + '.docx'
            docx_path = os.path.join(temp_dir, docx_filename)

            if not os.path.exists(docx_path):
                raise Exception("DOCX conversion failed - output file not found")

            # DOCX 파싱
            with open(docx_path, 'rb') as f:
                docx_bytes = f.read()

            result = self._parse_docx(docx_bytes)
            if result.success:
                result.method = "libreoffice"

            return result

    def is_valid_docx(self, file_bytes: bytes) -> bool:
        """DOCX 파일 유효성 체크"""
        if Document is None:
            return False

        try:
            Document(io.BytesIO(file_bytes))
            return True
        except Exception:
            return False
